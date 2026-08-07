from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx(meeting_data: dict, output_path: str):
    doc = Document()

    title = doc.add_heading(meeting_data["title"], level=1)

    paragraph = doc.add_paragraph(f"작성일: {meeting_data['created_at']}")
    paragraph.runs[0].font.size = Pt(10)
    paragraph.runs[0].font.color.rgb = RGBColor(0x7C, 0x76, 0x6A)

    summary = meeting_data.get("summary")

    if summary:
        doc.add_heading("요약", level=2)
        doc.add_paragraph(summary.get("summary_text", ""))

        decisions = summary.get("decisions", [])

        if decisions:
            doc.add_heading("결정사항", level=2)

            for decision in decisions:
                doc.add_paragraph(decision, style="List Bullet")

        todos = summary.get("todos", [])

        if todos:
            doc.add_heading("할 일", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"

            header_cells = table.rows[0].cells
            header_cells[0].text = "할 일"
            header_cells[1].text = "담당자"
            header_cells[2].text = "기한"

            for todo in todos:
                row_cells = table.add_row().cells
                row_cells[0].text = todo.get("task", "")
                row_cells[1].text = todo.get("assignee") or "미지정"
                row_cells[2].text = todo.get("due") or "미지정"

        next_agenda = summary.get("next_agenda", [])

        if next_agenda:
            doc.add_heading("다음 안건", level=2)

            for agenda in next_agenda:
                doc.add_paragraph(agenda, style="List Bullet")

    transcript = meeting_data.get("transcript")

    if transcript and transcript.get("segments"):
        doc.add_heading("전사문", level=2)

        for segment in transcript["segments"]:
            p = doc.add_paragraph()
            time_run = p.add_run(f"[{format_seconds(segment['start'])}]")
            time_run.font.bold = True
            time_run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
            p.add_run(segment["text"])

    doc.save(output_path)
    return output_path

def format_seconds(seconds):
    hour = int(seconds // 3600)
    minute = int((seconds % 3600) // 60)
    second = int(seconds % 60)

    return f"{hour:02d}:{minute:02d}:{second:02d}"